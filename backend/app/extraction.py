import re
import json
from io import BytesIO
from typing import Literal
from urllib.request import Request, urlopen

from google import genai
from google.genai import types
from pydantic import BaseModel, Field

from .settings import settings


class ExperienceExtraction(BaseModel):
    company: str | None
    role: str
    location: str | None
    start_date: str | None
    end_date: str | None
    is_current: bool
    description: str | None


class EducationExtraction(BaseModel):
    institution: str | None
    qualification: str
    field_of_study: str | None
    start_year: int | None
    end_year: int | None


class SkillExtraction(BaseModel):
    name: str
    category: str | None
    level: str | None


class CandidateExtraction(BaseModel):
    first_name: str
    last_name: str
    birth_year: int | None
    birth_place: str | None
    declared_gender: Literal["female", "male", "other"] | None
    email: str | None
    phone: str | None
    city: str | None
    bio: str | None
    experiences: list[ExperienceExtraction]
    education: list[EducationExtraction]
    skills: list[SkillExtraction]
    confidence: float = Field(ge=0, le=1)
    portrait_found: bool
    portrait_box_2d: list[int] = Field(description="Foto profilo nella prima pagina: [ymin,xmin,ymax,xmax], coordinate 0-1000")
    portrait_confidence: float = Field(ge=0, le=1)
    job_category: Literal["accounting", "logistics", "marketing", "cashier", "sales", "warehouse", "office", "other"]
    protected_category: bool


PROMPT = """Estrai dal curriculum i dati del candidato in modo fedele.
Non inventare informazioni mancanti. Non dedurre il sesso dal nome, dalla foto,
dalla nazionalità o da altri indizi: declared_gender va compilato solo se il CV
lo dichiara esplicitamente. Per le date usa YYYY-MM-DD quando il giorno è noto,
YYYY-MM-01 quando sono noti solo anno e mese, e YYYY-01-01 quando è noto solo
l'anno. La bio deve essere una sintesi professionale neutra di massimo 500
caratteri. Inserisci soltanto esperienze, istruzione e competenze presenti nel CV.
Se nella prima pagina è presente una chiara foto ritratto del candidato, imposta
portrait_found=true e portrait_box_2d con [ymin,xmin,ymax,xmax] normalizzati da 0
a 1000. Ignora loghi, icone, firme, gruppi e immagini decorative. Se il ritratto
non è chiaro usa false, lista vuota e confidenza 0. Classifica la mansione principale
considerando soprattutto l'esperienza più recente: accounting=contabilità, logistics=logistica
d'ufficio/coordinamento, marketing=marketing professionale, cashier=cassa, sales=vendita in
negozio, warehouse=magazzino operativo, office=segreteria/back office, other=altro.
Non usare una vecchia esperienza o una competenza marginale per prevalere sul ruolo recente.
protected_category può essere true soltanto se il CV dichiara esplicitamente l'appartenenza
alle categorie protette o la Legge 68/99; non dedurla da salute, foto o altri indizi.
"""


def _document_text(data: bytes, filename: str) -> str:
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if suffix == "pdf":
        import fitz
        document = fitz.open(stream=data, filetype="pdf")
        return "\n".join(page.get_text() for page in document)
    if suffix == "docx":
        from docx import Document
        document = Document(BytesIO(data))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
    raise ValueError("Il modello locale supporta PDF e DOCX; convertire i file DOC")


def _extract_candidate_ollama(data: bytes, filename: str) -> CandidateExtraction:
    text = _document_text(data, filename)
    if not text.strip():
        raise ValueError("Il curriculum non contiene testo leggibile")
    schema = CandidateExtraction.model_json_schema()
    payload = json.dumps({
        "model": settings.ollama_model, "stream": False, "format": schema,
        "messages": [{"role":"user","content":f"{PROMPT}\n\nRispondi secondo questo schema JSON:\n{json.dumps(schema)}\n\nCURRICULUM:\n{text[:24000]}"}],
        "options": {"temperature": 0, "num_ctx": 16384},
    }).encode()
    request = Request(f"{settings.ollama_base_url.rstrip('/')}/api/chat", data=payload, headers={"Content-Type":"application/json"})
    with urlopen(request, timeout=180) as response:
        result = json.loads(response.read())
    return CandidateExtraction.model_validate_json(result["message"]["content"])


def extract_candidate(data: bytes, filename: str) -> CandidateExtraction:
    if settings.ai_provider.lower() == "ollama":
        return _extract_candidate_ollama(data, filename)
    client = genai.Client(api_key=settings.gemini_api_key)
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if suffix == "pdf":
        content = [types.Part.from_bytes(data=data, mime_type="application/pdf"), PROMPT]
    elif suffix == "doc":
        content = [types.Part.from_bytes(data=data, mime_type="application/msword"), PROMPT]
    elif suffix == "docx":
        from docx import Document
        document = Document(BytesIO(data))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        content = [PROMPT, text]
    else:
        raise ValueError("Formato del curriculum non supportato")

    response = client.models.generate_content(
        model=settings.gemini_model,
        contents=content,
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
            response_schema=CandidateExtraction,
            temperature=0.1,
        ),
    )
    if response.parsed:
        return response.parsed
    return CandidateExtraction.model_validate_json(response.text)


def normalize_email(value: str | None) -> str | None:
    return value.strip().lower() if value and "@" in value else None


def normalize_phone(value: str | None) -> str | None:
    if not value:
        return None
    digits = re.sub(r"\D", "", value)
    return digits[-10:] if len(digits) >= 7 else None
